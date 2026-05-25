import os
import re
import json
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import List, Tuple, Optional
from urllib.parse import urlparse, parse_qs

import requests

from fastapi import FastAPI, Request, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from groq import Groq
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = FastAPI(title="Transcritor Groq Telegram Worker", version="1.1.0")

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
TRANSCRIPTION_MODEL = os.getenv("GROQ_TRANSCRIPTION_MODEL", "whisper-large-v3-turbo")
TRANSLATION_MODEL = os.getenv("GROQ_TRANSLATION_MODEL", "llama-3.3-70b-versatile")
MAX_AUDIO_MB = float(os.getenv("MAX_AUDIO_MB", "24"))
TARGET_AUDIO_BITRATE = os.getenv("TARGET_AUDIO_BITRATE", "24k")
TARGET_AUDIO_FORMAT = os.getenv("TARGET_AUDIO_FORMAT", "mp3")
TELEGRAM_BOT_TOKEN = os.getenv("TELEGRAM_BOT_TOKEN")

if GROQ_API_KEY:
    groq = Groq(api_key=GROQ_API_KEY)
else:
    groq = None


def safe_filename(name: str) -> str:
    name = name or "arquivo"
    name = re.sub(r"[^a-zA-Z0-9._-]+", "_", name)
    return name[:120]




def normalize_media_url(url: str) -> str:
    url = (url or "").strip()
    if not url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="URL inválida. Envie um link começando com http:// ou https://")

    # Dropbox: transforma link de visualização em download direto
    if "dropbox.com" in url:
        if "dl=0" in url:
            url = url.replace("dl=0", "dl=1")
        elif "dl=1" not in url:
            sep = "&" if "?" in url else "?"
            url = f"{url}{sep}dl=1"

    # Google Drive: tenta transformar link compartilhado em download direto
    if "drive.google.com" in url:
        parsed = urlparse(url)
        file_id = None
        if "/file/d/" in parsed.path:
            file_id = parsed.path.split("/file/d/")[1].split("/")[0]
        else:
            qs = parse_qs(parsed.query)
            file_id = (qs.get("id") or [None])[0]
        if file_id:
            url = f"https://drive.google.com/uc?export=download&id={file_id}"

    return url


def filename_from_url(url: str, fallback: str = "media") -> str:
    try:
        path = urlparse(url).path
        name = Path(path).name
        return safe_filename(name or fallback)
    except Exception:
        return safe_filename(fallback)


def download_url_to_file(url: str, dest: Path) -> None:
    url = normalize_media_url(url)
    session = requests.Session()
    with session.get(url, stream=True, timeout=(20, 180)) as r:
        r.raise_for_status()

        # Google Drive às vezes exige token de confirmação para arquivos grandes
        token = None
        for k, v in session.cookies.items():
            if k.startswith("download_warning"):
                token = v
                break

        if token:
            parsed = urlparse(url)
            qs = parse_qs(parsed.query)
            file_id = (qs.get("id") or [None])[0]
            if file_id:
                confirm_url = f"https://drive.google.com/uc?export=download&confirm={token}&id={file_id}"
                r.close()
                r = session.get(confirm_url, stream=True, timeout=(20, 180))
                r.raise_for_status()

        with dest.open("wb") as f:
            for chunk in r.iter_content(chunk_size=1024 * 1024):
                if chunk:
                    f.write(chunk)

        if dest.stat().st_size == 0:
            raise HTTPException(status_code=400, detail="O link foi baixado, mas o arquivo veio vazio.")


def download_telegram_file_id_to_file(file_id: str, dest: Path) -> None:
    if not TELEGRAM_BOT_TOKEN:
        raise HTTPException(status_code=500, detail="TELEGRAM_BOT_TOKEN não configurado no servidor.")
    if not file_id:
        raise HTTPException(status_code=400, detail="fileId obrigatório.")

    meta_url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/getFile"
    meta = requests.get(meta_url, params={"file_id": file_id}, timeout=30)
    meta.raise_for_status()
    data = meta.json()
    if not data.get("ok"):
        raise HTTPException(status_code=400, detail=f"Telegram não retornou o arquivo: {data}")

    file_path = data.get("result", {}).get("file_path")
    if not file_path:
        raise HTTPException(status_code=400, detail="Telegram não retornou file_path para esse arquivo.")

    download_url = f"https://api.telegram.org/file/bot{TELEGRAM_BOT_TOKEN}/{file_path}"
    with requests.get(download_url, stream=True, timeout=(20, 600)) as r:
        r.raise_for_status()
        with dest.open("wb") as f:
            for chunk in r.iter_content(chunk_size=1024 * 1024):
                if chunk:
                    f.write(chunk)

    if dest.stat().st_size == 0:
        raise HTTPException(status_code=400, detail="O arquivo do Telegram foi baixado vazio.")

def run(cmd: List[str]) -> None:
    result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"Comando falhou: {' '.join(cmd)}\n{result.stderr[-2000:]}")


def file_mb(path: Path) -> float:
    return path.stat().st_size / (1024 * 1024)


def convert_to_small_audio(input_path: Path, output_path: Path) -> None:
    # Mono, 16 kHz, baixa taxa. Bom para fala e tenta ficar abaixo do limite.
    run([
        "ffmpeg", "-y", "-i", str(input_path),
        "-vn", "-ac", "1", "-ar", "16000",
        "-b:a", TARGET_AUDIO_BITRATE,
        str(output_path)
    ])


def split_audio(input_audio: Path, chunks_dir: Path) -> List[Path]:
    chunks_dir.mkdir(parents=True, exist_ok=True)
    pattern = chunks_dir / f"chunk_%03d.{TARGET_AUDIO_FORMAT}"
    # 10 minutos por chunk. Reencoda para garantir tamanho baixo por parte.
    run([
        "ffmpeg", "-y", "-i", str(input_audio),
        "-f", "segment", "-segment_time", "600",
        "-ac", "1", "-ar", "16000", "-b:a", TARGET_AUDIO_BITRATE,
        str(pattern)
    ])
    return sorted(chunks_dir.glob(f"*.{TARGET_AUDIO_FORMAT}"))


def transcribe_one(audio_path: Path) -> dict:
    if groq is None:
        raise RuntimeError("GROQ_API_KEY não configurada.")
    with audio_path.open("rb") as f:
        result = groq.audio.transcriptions.create(
            file=(audio_path.name, f.read()),
            model=TRANSCRIPTION_MODEL,
            response_format="verbose_json",
            temperature=0,
        )
    # O SDK pode retornar objeto pydantic ou dict dependendo da versão
    if hasattr(result, "model_dump"):
        return result.model_dump()
    if isinstance(result, dict):
        return result
    return json.loads(result.json())


def transcribe_audio(audio_path: Path, work_dir: Path) -> Tuple[str, str, List[dict]]:
    # Tenta arquivo único. Se passar do limite, divide automaticamente.
    segments: List[dict] = []
    language = "unknown"

    if file_mb(audio_path) <= MAX_AUDIO_MB:
        data = transcribe_one(audio_path)
        text = data.get("text", "")
        language = data.get("language") or language
        segments = data.get("segments") or []
        if not segments and text:
            segments = [{"start": 0, "end": 0, "text": text}]
        return text, language, segments

    chunks = split_audio(audio_path, work_dir / "chunks")
    full_text_parts = []
    offset = 0.0
    for chunk in chunks:
        data = transcribe_one(chunk)
        language = data.get("language") or language
        text = data.get("text", "")
        full_text_parts.append(text)
        chunk_segments = data.get("segments") or []
        for seg in chunk_segments:
            start = float(seg.get("start", 0)) + offset
            end = float(seg.get("end", 0)) + offset
            segments.append({"start": start, "end": end, "text": seg.get("text", "")})
        # Duração aproximada via ffprobe
        try:
            probe = subprocess.run([
                "ffprobe", "-v", "error", "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1", str(chunk)
            ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
            offset += float(probe.stdout.strip() or 600)
        except Exception:
            offset += 600
    return "\n".join(full_text_parts), language, segments


def needs_translation(language: str) -> bool:
    if not language:
        return True
    lang = language.lower()
    return not (lang.startswith("pt") or "portugu" in lang)


def chunk_text(text: str, max_chars: int = 9000) -> List[str]:
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks, current = [], ""
    for p in paragraphs:
        if len(current) + len(p) + 2 > max_chars and current:
            chunks.append(current)
            current = p
        else:
            current = f"{current}\n\n{p}" if current else p
    if current:
        chunks.append(current)
    return chunks or [text]


def translate_to_ptbr(text: str, language: str) -> str:
    if groq is None:
        raise RuntimeError("GROQ_API_KEY não configurada.")
    if not needs_translation(language):
        return text

    translated_parts = []
    for part in chunk_text(text):
        completion = groq.chat.completions.create(
            model=TRANSLATION_MODEL,
            temperature=0,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "Você é um tradutor profissional de copy, anúncios e VSLs. "
                        "Traduza para português brasileiro com máxima fidelidade. "
                        "Não resuma, não melhore, não censure e não adicione ideias. "
                        "Preserve estrutura, repetições, tom emocional, ganchos, CTAs e quebras de parágrafo. "
                        "Quando houver expressão idiomática, traduza pelo sentido natural em PT-BR."
                    ),
                },
                {
                    "role": "user",
                    "content": f"Idioma original detectado: {language}\n\nTexto original:\n{part}",
                },
            ],
        )
        translated_parts.append(completion.choices[0].message.content.strip())
    return "\n\n".join(translated_parts)


def fmt_time(seconds: float) -> str:
    seconds = max(0, int(seconds or 0))
    h = seconds // 3600
    m = (seconds % 3600) // 60
    s = seconds % 60
    return f"{h:02d}:{m:02d}:{s:02d}"


def make_docx(out_path: Path, original_name: str, language: str, translated_text: str, original_text: str, segments: List[dict]) -> None:
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("Transcrição / Tradução PT-BR", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.add_run("Arquivo: ").bold = True
    p.add_run(original_name)
    p = doc.add_paragraph()
    p.add_run("Idioma original detectado: ").bold = True
    p.add_run(language or "não identificado")

    doc.add_paragraph("")
    doc.add_heading("Versão em português brasileiro", level=1)

    # Se houver segmentos, usa os timestamps como guia. Se foi traduzido em bloco, coloca texto corrido.
    if translated_text.strip():
        for block in [b.strip() for b in translated_text.split("\n\n") if b.strip()]:
            doc.add_paragraph(block)

    doc.add_page_break()
    doc.add_heading("Transcrição original", level=1)
    if segments:
        for seg in segments:
            start = fmt_time(seg.get("start", 0))
            end = fmt_time(seg.get("end", 0))
            p = doc.add_paragraph()
            run_time = p.add_run(f"[{start} – {end}] ")
            run_time.bold = True
            p.add_run((seg.get("text") or "").strip())
    else:
        doc.add_paragraph(original_text)

    doc.save(out_path)


@app.get("/health")
def health():
    return {"ok": True, "groq_key_configured": bool(GROQ_API_KEY), "telegram_token_configured": bool(TELEGRAM_BOT_TOKEN)}



@app.post("/process-source")
async def process_source(request: Request):
    if groq is None:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY não configurada no servidor.")

    try:
        payload = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Envie JSON com sourceType.")

    source_type = payload.get("sourceType")
    original_name = safe_filename(payload.get("fileName") or "media")
    suffix = Path(original_name).suffix or ".mp4"

    with tempfile.TemporaryDirectory() as tmp:
        work_dir = Path(tmp)
        input_path = work_dir / f"input{suffix}"
        audio_path = work_dir / f"audio.{TARGET_AUDIO_FORMAT}"
        output_docx = work_dir / f"{Path(original_name).stem}_transcricao_ptbr.docx"

        if source_type == "url":
            media_url = payload.get("url") or payload.get("media_url")
            if not media_url:
                raise HTTPException(status_code=400, detail="URL obrigatória para sourceType=url.")
            download_url_to_file(media_url, input_path)
        elif source_type == "telegram_file_id":
            file_id = payload.get("fileId")
            download_telegram_file_id_to_file(file_id, input_path)
        else:
            raise HTTPException(status_code=400, detail="sourceType inválido. Use url ou telegram_file_id.")

        convert_to_small_audio(input_path, audio_path)
        original_text, language, segments = transcribe_audio(audio_path, work_dir)
        translated_text = translate_to_ptbr(original_text, language)
        make_docx(output_docx, original_name, language, translated_text, original_text, segments)

        final_path = Path(tempfile.gettempdir()) / output_docx.name
        shutil.copy2(output_docx, final_path)

    return FileResponse(
        path=str(final_path),
        filename=final_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

@app.post("/process-url-media")
async def process_url_media(request: Request):
    if groq is None:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY não configurada no servidor.")

    try:
        payload = await request.json()
    except Exception:
        raise HTTPException(status_code=400, detail="Envie JSON com media_url.")

    media_url = payload.get("media_url") or payload.get("url")
    if not media_url:
        raise HTTPException(status_code=400, detail="Campo media_url obrigatório.")

    original_name = safe_filename(payload.get("file_name") or filename_from_url(media_url, "media"))
    suffix = Path(original_name).suffix or ".mp4"

    with tempfile.TemporaryDirectory() as tmp:
        work_dir = Path(tmp)
        input_path = work_dir / f"input{suffix}"
        audio_path = work_dir / f"audio.{TARGET_AUDIO_FORMAT}"
        output_docx = work_dir / f"{Path(original_name).stem}_transcricao_ptbr.docx"

        download_url_to_file(media_url, input_path)
        convert_to_small_audio(input_path, audio_path)
        original_text, language, segments = transcribe_audio(audio_path, work_dir)
        translated_text = translate_to_ptbr(original_text, language)
        make_docx(output_docx, original_name, language, translated_text, original_text, segments)

        final_path = Path(tempfile.gettempdir()) / output_docx.name
        shutil.copy2(output_docx, final_path)

    return FileResponse(
        path=str(final_path),
        filename=final_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.post("/process-telegram-media")
async def process_telegram_media(request: Request):
    if groq is None:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY não configurada no servidor.")
    content = await request.body()
    if not content:
        raise HTTPException(status_code=400, detail="Nenhum arquivo recebido.")

    original_name = safe_filename(request.headers.get("x-file-name", "telegram_media"))
    suffix = Path(original_name).suffix or ".bin"

    with tempfile.TemporaryDirectory() as tmp:
        work_dir = Path(tmp)
        input_path = work_dir / f"input{suffix}"
        audio_path = work_dir / f"audio.{TARGET_AUDIO_FORMAT}"
        output_docx = work_dir / f"{Path(original_name).stem}_transcricao_ptbr.docx"

        input_path.write_bytes(content)
        convert_to_small_audio(input_path, audio_path)
        original_text, language, segments = transcribe_audio(audio_path, work_dir)
        translated_text = translate_to_ptbr(original_text, language)
        make_docx(output_docx, original_name, language, translated_text, original_text, segments)

        # Copia para um caminho fora do TemporaryDirectory antes de responder
        final_path = Path(tempfile.gettempdir()) / output_docx.name
        shutil.copy2(output_docx, final_path)

    return FileResponse(
        path=str(final_path),
        filename=final_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
